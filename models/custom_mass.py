"""
Model for complete Mass structure with customization capabilities
"""

from dataclasses import dataclass, field
from typing import Optional, List, Dict
from datetime import date
from .base import Reading, Psalm, Prayer, Antiphon, Celebration, LiturgicalColor


@dataclass
class MassPart:
    """Represents a part of the Mass"""
    title: str
    content: str
    order: int = 0
    
    def __str__(self):
        return f"\n{self.title}\n{'=' * len(self.title)}\n{self.content}\n"


class CustomMass:
    """
    Represents a fully customizable Catholic Mass with all its parts.
    
    The Mass is structured in four main sections:
    1. Ritos Iniciais (Introductory Rites)
    2. Liturgia da Palavra (Liturgy of the Word)
    3. Liturgia Eucarística (Liturgy of the Eucharist)
    4. Ritos Finais (Concluding Rites)
    """
    
    def __init__(self):
        self.celebration: Optional[Celebration] = None
        self.parts: Dict[str, MassPart] = {}
        self._init_default_structure()
    
    def _init_default_structure(self):
        """Initialize with complete Mass structure including all possible parts"""
        # Ritos Iniciais
        self.parts["entrance_procession"] = MassPart("Procissão de Entrada", "", 1)
        self.parts["entrance_antiphon"] = MassPart("Antífona de Entrada", "", 2)
        self.parts["entrance_hymn"] = MassPart("Canto de Entrada", "", 3)
        self.parts["sign_cross"] = MassPart("Sinal da Cruz", "", 4)
        self.parts["greeting"] = MassPart("Saudação", "Em nome do Pai, do Filho e do Espírito Santo.", 5)
        self.parts["introduction"] = MassPart("Introdução à Celebração", "", 6)
        self.parts["blessing_water"] = MassPart("Bênção e Aspersão da Água (opcional)", "", 7)
        self.parts["penitential_intro"] = MassPart("Introdução ao Ato Penitencial", "", 8)
        self.parts["penitential"] = MassPart("Ato Penitencial", "", 9)
        self.parts["kyrie"] = MassPart("Kyrie", "Senhor, tende piedade de nós.", 10)
        self.parts["gloria"] = MassPart("Glória", "", 11)
        self.parts["collect"] = MassPart("Oração do Dia (Coleta)", "", 12)
        
        # Liturgia da Palavra
        self.parts["first_reading"] = MassPart("Primeira Leitura", "", 13)
        self.parts["responsorial_gradual"] = MassPart("Canto Responsorial/Gradual", "", 14)
        self.parts["psalm"] = MassPart("Salmo Responsorial", "", 15)
        self.parts["second_reading"] = MassPart("Segunda Leitura", "", 16)
        self.parts["sequence"] = MassPart("Sequência (dias especiais)", "", 17)
        self.parts["gospel_acclamation"] = MassPart("Aclamação ao Evangelho (Aleluia)", "", 18)
        self.parts["gospel_procession"] = MassPart("Procissão do Evangelho", "", 19)
        self.parts["gospel"] = MassPart("Evangelho", "", 20)
        self.parts["homily"] = MassPart("Homilia", "", 21)
        self.parts["silence_reflection"] = MassPart("Silêncio para Reflexão", "", 22)
        self.parts["creed"] = MassPart("Profissão de Fé (Credo)", "", 23)
        self.parts["prayers_faithful_intro"] = MassPart("Introdução à Oração dos Fiéis", "", 24)
        self.parts["prayers_faithful"] = MassPart("Oração dos Fiéis", "", 25)
        self.parts["prayers_faithful_conclusion"] = MassPart("Conclusão da Oração dos Fiéis", "", 26)
        
        # Liturgia Eucarística
        # Preparação das Oferendas
        self.parts["offertory_procession"] = MassPart("Procissão das Oferendas", "", 27)
        self.parts["offertory_hymn"] = MassPart("Canto das Oferendas", "", 28)
        self.parts["offertory"] = MassPart("Apresentação das Oferendas", "", 29)
        self.parts["preparation_bread"] = MassPart("Preparação do Pão", "", 30)
        self.parts["preparation_wine"] = MassPart("Preparação do Vinho", "", 31)
        self.parts["mixing_water_wine"] = MassPart("Mistura da Água com o Vinho", "", 32)
        self.parts["offering_incense"] = MassPart("Incensação das Oferendas (opcional)", "", 33)
        self.parts["washing_hands"] = MassPart("Lavabo (Lavagem das Mãos)", "", 34)
        self.parts["invitation_prayer"] = MassPart("Convite à Oração", "", 35)
        self.parts["prayer_offerings"] = MassPart("Oração sobre as Oferendas", "", 36)
        
        # Oração Eucarística
        self.parts["preface_dialogue"] = MassPart("Diálogo do Prefácio", "", 37)
        self.parts["preface"] = MassPart("Prefácio", "", 38)
        self.parts["sanctus"] = MassPart("Santo", "", 39)
        self.parts["epiclesis_1"] = MassPart("Primeira Epiclese (invocação do Espírito Santo)", "", 40)
        self.parts["institution_narrative"] = MassPart("Narrativa da Instituição", "", 41)
        self.parts["consecration_bread"] = MassPart("Consagração do Pão", "", 42)
        self.parts["consecration_wine"] = MassPart("Consagração do Vinho", "", 43)
        self.parts["elevation"] = MassPart("Elevação", "", 44)
        self.parts["mystery_faith"] = MassPart("Mistério da Fé (Aclamação)", "", 45)
        self.parts["anamnesis"] = MassPart("Anamnese (Memorial)", "", 46)
        self.parts["epiclesis_2"] = MassPart("Segunda Epiclese", "", 47)
        self.parts["intercessions"] = MassPart("Intercessões", "", 48)
        self.parts["doxology"] = MassPart("Doxologia Final", "", 49)
        self.parts["great_amen"] = MassPart("Grande Amém", "", 50)
        
        # Rito da Comunhão
        self.parts["our_father_intro"] = MassPart("Introdução ao Pai Nosso", "", 51)
        self.parts["our_father"] = MassPart("Pai Nosso", "", 52)
        self.parts["embolism"] = MassPart("Embolismo (Livrai-nos de todos os males)", "", 53)
        self.parts["doxology_prayer"] = MassPart("Doxologia do Povo", "", 54)
        self.parts["peace_prayer"] = MassPart("Oração pela Paz", "", 55)
        self.parts["peace"] = MassPart("Rito da Paz", "", 56)
        self.parts["peace_exchange"] = MassPart("Saudação da Paz", "", 57)
        self.parts["fraction"] = MassPart("Fração do Pão", "", 58)
        self.parts["agnus_dei"] = MassPart("Cordeiro de Deus", "", 59)
        self.parts["commingling"] = MassPart("Imissão (mistura do pão e vinho)", "", 60)
        self.parts["private_preparation"] = MassPart("Oração Privada do Sacerdote", "", 61)
        self.parts["invitation_communion"] = MassPart("Convite à Comunhão", "", 62)
        self.parts["communion"] = MassPart("Comunhão", "", 63)
        self.parts["communion_antiphon"] = MassPart("Antífona da Comunhão", "", 64)
        self.parts["communion_hymn"] = MassPart("Canto de Comunhão", "", 65)
        self.parts["communion_meditation"] = MassPart("Momento de Ação de Graças", "", 66)
        self.parts["purification"] = MassPart("Purificação dos Vasos Sagrados", "", 67)
        self.parts["silence_thanksgiving"] = MassPart("Silêncio para Ação de Graças", "", 68)
        self.parts["prayer_communion"] = MassPart("Oração depois da Comunhão", "", 69)
        
        # Ritos Finais
        self.parts["greeting_final"] = MassPart("Saudação Final", "", 70)
        self.parts["announcements"] = MassPart("Avisos", "", 71)
        self.parts["blessing_introduction"] = MassPart("Introdução à Bênção", "", 72)
        self.parts["solemn_blessing"] = MassPart("Bênção Solene (opcional)", "", 73)
        self.parts["blessing"] = MassPart("Bênção", "", 74)
        self.parts["dismissal"] = MassPart("Despedida", "", 75)
        self.parts["recession"] = MassPart("Procissão de Saída", "", 76)
        self.parts["final_hymn"] = MassPart("Canto Final", "", 77)
    
    def set_celebration(self, name: str, date_str: Optional[str] = None, 
                       celebration_type: str = "solenidade",
                       color: str = "branco", season: str = "tempo comum"):
        """Set the celebration details"""
        cel_date = date.today()
        if date_str:
            from datetime import datetime
            cel_date = datetime.strptime(date_str, "%Y-%m-%d").date()
        
        self.celebration = Celebration(
            name=name,
            date=cel_date,
            type=celebration_type,
            color=LiturgicalColor(color),
            season=season
        )
    
    def set_entrance_antiphon(self, text: str, reference: str = ""):
        """Set the entrance antiphon"""
        antiphon = Antiphon("Entrada", text, reference)
        self.parts["entrance_antiphon"].content = str(antiphon)
    
    def set_communion_antiphon(self, text: str, reference: str = ""):
        """Set the communion antiphon"""
        antiphon = Antiphon("Comunhão", text, reference)
        self.parts["communion_antiphon"].content = str(antiphon)
    
    def set_readings(self, first_reading: str = "", psalm: str = "", 
                    second_reading: str = "", gospel: str = ""):
        """Set the readings for the Mass"""
        if first_reading:
            self.parts["first_reading"].content = f"Leitura: {first_reading}"
        if psalm:
            self.parts["psalm"].content = f"Salmo: {psalm}"
        if second_reading:
            self.parts["second_reading"].content = f"Leitura: {second_reading}"
        if gospel:
            self.parts["gospel"].content = f"Evangelho: {gospel}"
    
    def set_part_content(self, part_key: str, content: str):
        """Set content for a specific part of the Mass"""
        if part_key in self.parts:
            self.parts[part_key].content = content
    
    def add_custom_prayer(self, title: str, text: str, position: int = 14):
        """Add a custom prayer at a specific position"""
        # Generate a unique key by appending a counter if needed
        base_key = title.lower().replace(" ", "_")
        key = base_key
        counter = 1
        while key in self.parts:
            key = f"{base_key}_{counter}"
            counter += 1
        self.parts[key] = MassPart(title, text, position)
    
    def _get_sorted_parts(self):
        """Get parts sorted by order"""
        return sorted(self.parts.values(), key=lambda x: x.order)
    
    def get_full_text(self) -> str:
        """Get the complete formatted text of the Mass"""
        result = []
        
        if self.celebration:
            result.append(f"\n{'=' * 80}")
            result.append(f"{self.celebration.name.upper()}")
            result.append(f"Data: {self.celebration.date.strftime('%d/%m/%Y')}")
            result.append(f"Cor Litúrgica: {self.celebration.color}")
            result.append(f"{'=' * 80}\n")
        
        # Use helper method to get sorted parts
        for part in self._get_sorted_parts():
            if part.content:  # Only include parts with content
                result.append(str(part))
        
        return "\n".join(result)
    
    def export_to_text(self, filename: str):
        """Export the Mass to a text file"""
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(self.get_full_text())
    
    def export_to_pdf(self, filename: str, **options):
        """
        Export the Mass to PDF format with customization options
        
        Options:
            font_family: str - Font family (Times-Roman, Helvetica, Courier)
            font_size: int - Base font size (default: 12)
            page_size: str - Page size (A4, Letter, A5)
            margins: int - Margin size in points (default: 72)
            title_size: int - Title font size (default: 18)
            include_header: bool - Include header (default: True)
            include_footer: bool - Include footer (default: True)
            liturgical_color: str - Liturgical color for accent (default: verde)
        """
        try:
            from reportlab.lib.pagesizes import letter, A4, A5
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            
            # Get options with defaults
            font_family = options.get('font_family', 'Times-Roman')
            font_size = options.get('font_size', 12)
            page_size_str = options.get('page_size', 'A4')
            margins = options.get('margins', 72)
            title_size = options.get('title_size', 18)
            include_header = options.get('include_header', True)
            include_footer = options.get('include_footer', True)
            
            # Map page size string to actual page size
            page_sizes = {
                'A4': A4,
                'Letter': letter,
                'A5': A5
            }
            page_size = page_sizes.get(page_size_str, A4)
            
            # Create document
            doc = SimpleDocTemplate(
                filename, 
                pagesize=page_size,
                leftMargin=margins,
                rightMargin=margins,
                topMargin=margins,
                bottomMargin=margins
            )
            
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=font_family,
                fontSize=title_size,
                textColor=colors.black,
                spaceAfter=20,
                alignment=1,  # Center
                leading=title_size * 1.2
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontName=font_family,
                fontSize=font_size + 2,
                textColor=colors.HexColor('#5e72e4'),
                spaceAfter=10,
                spaceBefore=15,
                leading=(font_size + 2) * 1.3
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontName=font_family,
                fontSize=font_size,
                textColor=colors.black,
                spaceAfter=6,
                leading=font_size * 1.4,
                alignment=0  # Left
            )
            
            # Add header if requested
            if include_header and self.celebration:
                story.append(Paragraph(self.celebration.name.upper(), title_style))
                info_text = f"Data: {self.celebration.date.strftime('%d/%m/%Y')}"
                if hasattr(self.celebration, 'color') and self.celebration.color:
                    color_name = str(self.celebration.color).title()
                    info_text += f" | Cor Litúrgica: {color_name}"
                story.append(Paragraph(info_text, body_style))
                story.append(Spacer(1, 0.3*inch))
            
            # Add content
            for part in self._get_sorted_parts():
                if part.content:
                    # Add part title
                    story.append(Paragraph(f"<b>{part.title}</b>", heading_style))
                    
                    # Add part content
                    content = part.content.replace('\n', '<br/>')
                    story.append(Paragraph(content, body_style))
                    story.append(Spacer(1, 0.15*inch))
            
            # Add footer if requested
            if include_footer:
                story.append(Spacer(1, 0.5*inch))
                footer_style = ParagraphStyle(
                    'Footer',
                    parent=body_style,
                    fontSize=font_size - 2,
                    textColor=colors.grey,
                    alignment=1
                )
                story.append(Paragraph("Folheto de Missa - Liturgia Católica", footer_style))
            
            # Build PDF
            doc.build(story)
            
        except ImportError:
            raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")
    
    def export_to_docx(self, filename: str):
        """Export the Mass to DOCX format"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Add title
            if self.celebration:
                title = doc.add_heading(self.celebration.name.upper(), 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                date_para = doc.add_paragraph(f"Data: {self.celebration.date.strftime('%d/%m/%Y')}")
                date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
            
            # Add content
            for part in self._get_sorted_parts():
                if part.content:
                    doc.add_heading(part.title, level=2)
                    doc.add_paragraph(part.content)
            
            doc.save(filename)
        except ImportError:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
